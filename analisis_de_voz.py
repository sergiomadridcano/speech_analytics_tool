import os
import sys
import time
import wave
import shutil
import tempfile
import subprocess
import tkinter as tk
from tkinter import filedialog
from datetime import datetime
from docx import Document
from faster_whisper import WhisperModel

# === CONFIG ===
EXTENSIONES_VALIDAS = {".mp3", ".wav", ".gsm", ".aac", ".mpeg", ".mpg", ".m4a", ".mp4"}

WHISPER_MODEL_SIZE = "small"
DEVICE = "cpu"
COMPUTE_TYPE = "int8_float32"   # si quieres aún más rápido: "int8"

# Speed/quality knobs (orientado a velocidad)
USE_VAD = True
BEAM_SIZE = 2   # más bajo => más rápido pero peor calidad
TEMPERATURE = 0.0

# VAD tuning (agresivo para llamadas; ajustable)
VAD_PARAMS = {
    "min_silence_duration_ms": 500,   # más bajo => corta más silencios
    "speech_pad_ms": 150,             # padding alrededor del speech
}

# Threads internos del motor (importante)
CPU_THREADS = max(1, (os.cpu_count() or 4) - 1)
NUM_WORKERS = 2  # prueba 2; si tu CPU es débil, deja 1

# === Utils ===
def fmt_hhmmss(segundos: float) -> str:
    if not segundos or segundos < 0:
        segundos = 0
    segundos = int(round(segundos))
    h = segundos // 3600
    m = (segundos % 3600) // 60
    s = segundos % 60
    return f"{h:02d}:{m:02d}:{s:02d}"

def seleccionar_carpeta():
    root = tk.Tk()
    root.withdraw()
    return filedialog.askdirectory(title="Selecciona la carpeta con audios")

def ensure_ffmpeg():
    if shutil.which("ffmpeg") is None:
        raise RuntimeError("No encuentro ffmpeg en PATH. Instálalo o agrégalo al PATH.")

def wav_es_16k_mono_pcm(wav_path: str) -> bool:
    # Validación ultra rápida sin decodificar todo
    try:
        with wave.open(wav_path, "rb") as wf:
            return (
                wf.getnchannels() == 1 and
                wf.getframerate() == 16000 and
                wf.getsampwidth() == 2   # 16-bit PCM
            )
    except Exception:
        return False

def wav_duration_seconds(wav_path: str) -> float | None:
    try:
        with wave.open(wav_path, "rb") as wf:
            frames = wf.getnframes()
            rate = wf.getframerate()
            return round(frames / float(rate), 2) if rate else None
    except Exception:
        return None

def convertir_con_ffmpeg(input_path: str, out_wav_path: str) -> None:
    # Conversión determinista y rápida
    # -vn: sin video; -ac 1 mono; -ar 16000; pcm_s16le 16-bit PCM
    cmd = [
        "ffmpeg", "-y",
        "-hide_banner", "-loglevel", "error",
        "-i", input_path,
        "-vn",
        "-ac", "1",
        "-ar", "16000",
        "-c:a", "pcm_s16le",
        out_wav_path
    ]
    subprocess.run(cmd, check=True)

def cargar_modelo_whisper():
    print(f"🧠 Cargando Whisper ({WHISPER_MODEL_SIZE}) | device={DEVICE} compute={COMPUTE_TYPE} | threads={CPU_THREADS} workers={NUM_WORKERS}")
    t0 = time.time()
    model = WhisperModel(
        WHISPER_MODEL_SIZE,
        device=DEVICE,
        compute_type=COMPUTE_TYPE,
        cpu_threads=CPU_THREADS,
        num_workers=NUM_WORKERS,
    )
    print(f"✅ Modelo Whisper cargado en {round(time.time() - t0, 2)}s")
    return model

def transcribir(model: WhisperModel, wav_path: str):
    t0 = time.time()

    kwargs = dict(
        language="es",
        vad_filter=USE_VAD,
        beam_size=BEAM_SIZE,
        temperature=TEMPERATURE,
    )
    if USE_VAD:
        kwargs["vad_parameters"] = VAD_PARAMS

    segments, info = model.transcribe(wav_path, **kwargs)
    texto = " ".join(seg.text.strip() for seg in segments if seg.text).strip()

    return texto, round(time.time() - t0, 2)

def listar_archivos(carpeta: str):
    # Ignora basura: temporales o convertidos previos
    out = []
    for f in os.listdir(carpeta):
        if f.startswith("."):
            continue
        ext = os.path.splitext(f)[1].lower()
        if ext not in EXTENSIONES_VALIDAS:
            continue
        if f.lower().endswith("_convertido.wav"):
            continue
        if f.lower().startswith("__temp__"):
            continue
        out.append(f)
    out.sort()
    return out

def analizar_carpeta(carpeta: str):
    ensure_ffmpeg()

    archivos = listar_archivos(carpeta)
    total = len(archivos)
    if total == 0:
        print("❌ No se encontraron archivos de audio.")
        return

    model = cargar_modelo_whisper()

    # Directorio temporal (no ensucia la carpeta del usuario)
    tmp_dir = tempfile.mkdtemp(prefix="whisper_tmp_")

    doc = Document()
    doc.add_heading("Transcripción de Audios (Whisper)", level=1)
    doc.add_paragraph(f"Carpeta analizada: {carpeta}")
    doc.add_paragraph(f"Total archivos: {total}")
    doc.add_paragraph(f"Modelo: {WHISPER_MODEL_SIZE} | device={DEVICE} | compute={COMPUTE_TYPE}")
    doc.add_paragraph(f"VAD: {USE_VAD} | beam_size: {BEAM_SIZE} | temp: {TEMPERATURE}")
    doc.add_paragraph(f"Threads: {CPU_THREADS} | Workers: {NUM_WORKERS}")
    doc.add_paragraph("")

    t0_global = time.time()
    tiempos = []

    try:
        for idx, archivo in enumerate(archivos, start=1):
            porcentaje = round((idx - 1) / total * 100, 1)
            eta = None
            if tiempos:
                avg = sum(tiempos) / len(tiempos)
                eta = avg * (total - (idx - 1))
            eta_str = fmt_hhmmss(eta) if eta else "--:--:--"

            print(f"📊 {idx}/{total} | {porcentaje}% | ETA {eta_str} | Iniciando: {archivo}")
            t0_file = time.time()

            in_path = os.path.join(carpeta, archivo)
            base = os.path.splitext(os.path.basename(archivo))[0]
            wav_path = os.path.join(tmp_dir, f"{base}.wav")

            # Fast-path: si ya es WAV 16k mono PCM, úsalo directo
            t_conv0 = time.time()
            if os.path.splitext(in_path)[1].lower() == ".wav" and wav_es_16k_mono_pcm(in_path):
                # copiamos al tmp para evitar locks / rutas raras y trabajar homogéneo
                shutil.copy2(in_path, wav_path)
            else:
                convertir_con_ffmpeg(in_path, wav_path)
            t_conv = round(time.time() - t_conv0, 2)

            dur = wav_duration_seconds(wav_path)

            texto, t_tr = transcribir(model, wav_path)

            # Doc
            doc.add_heading(f"Archivo: {archivo}", level=2)
            doc.add_paragraph(f"📌 Duración: {dur if dur is not None else 'Desconocida'} s")
            doc.add_paragraph(f"🔄 Conversión: {t_conv} s")
            doc.add_paragraph(f"⏱️ Transcripción: {t_tr} s")
            doc.add_paragraph(texto if texto else "[Sin texto detectado]")
            doc.add_paragraph("")

            # Limpieza por archivo (mantén tmp liviano)
            try:
                os.remove(wav_path)
            except Exception:
                pass

            t_total = time.time() - t0_file
            tiempos.append(t_total)

            pct_post = round(idx / total * 100, 1)
            avg = sum(tiempos) / len(tiempos)
            eta_post = avg * (total - idx)
            print(f"✅ Listo: {archivo} | {pct_post}% completado | ETA {fmt_hhmmss(eta_post)} | t={fmt_hhmmss(t_total)}")

    finally:
        # Limpieza temp dir completa
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass

    fecha = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    out_docx = os.path.join(carpeta, f"{fecha}_transcripcion_whisper.docx")
    doc.save(out_docx)

    print(f"\n🏁 Terminado. Tiempo total: {fmt_hhmmss(time.time() - t0_global)}")
    print(f"✅ Transcripción guardada en: {out_docx}")

if __name__ == "__main__":
    carpeta = seleccionar_carpeta()
    if carpeta:
        analizar_carpeta(carpeta)
    else:
        print("⚠️ No se seleccionó ninguna carpeta.")
